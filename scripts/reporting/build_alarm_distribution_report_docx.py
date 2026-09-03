"""Build the formal alarm-distribution validation report as a styled DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "reports" / "univariate_multivariate_distribution_adaptation_analysis_2026-09-03.md"
OUTPUT = ROOT / "docs" / "reports" / "IIA_univariate_multivariate_distribution_adaptation_report_2026-09-03.docx"

BLUE = "2E74B5"
NAVY = "17365D"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "DDEBF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
TEXT = "222222"

FIGURES = {
    "5.1 单变量分布": (
        ROOT / "experiments" / "paper_harness" / "univariate_adaptation" / "Figure_1.png",
        "图 1　单变量数据集分布漂移审计（测试统计仅用于冻结后的解释）",
    ),
    "5.2 多变量分布": (
        ROOT / "experiments" / "paper_harness" / "multivariate_adaptation" / "Figure_3.png",
        "图 2　多变量正常 train→evaluation 分布与相关结构漂移",
    ),
    "6.2 分布适配消融": (
        ROOT / "experiments" / "paper_harness" / "univariate_adaptation" / "Figure_3.png",
        "图 3　单变量适配消融：F1、FAR、MAR 与选择性覆盖率",
    ),
    "7.2 多变量适配前后": (
        ROOT / "experiments" / "paper_harness" / "multivariate_adaptation" / "Figure_1.png",
        "图 4　多变量 M0–M3 的 F1 与 coverage；TEP M3 为 67% 选择性结果",
    ),
}

PAGE_BREAK_BEFORE: set[str] = set()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "B7C9DD")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def new_numbering_instance(doc: Document, style_name: str = "List Number") -> int:
    """Clone a list numbering instance with an explicit level-0 restart."""
    style = doc.styles[style_name]
    num_pr = style._element.pPr.numPr
    base_num_id = int(num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(
        item for item in numbering.findall(qn("w:num")) if int(item.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    existing = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    new_num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.set(qn("w:val"), "0")
    num = num_pr.get_or_add_numId()
    num.set(qn("w:val"), str(num_id))


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str, size: float | None = None) -> None:
    """Add a small, safe subset of Markdown inline formatting."""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for token in token_re.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt((size or 10.5) - 0.3)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("IIA Benchmark · 内部验证报告　")
    set_run_font(run, size=8.5, color="667085")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    page_run = paragraph.add_run()
    page_run._r.append(fld_char1)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char2)
    set_run_font(page_run, size=8.5, color="667085")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string("52606D")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "IIA BENCHMARK  /  ALARM ANALYTICS"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=8.5, bold=True, color="667085")

    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, lines: list[str]) -> int:
    doc.add_paragraph()
    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strap.paragraph_format.space_before = Pt(44)
    strap.paragraph_format.space_after = Pt(18)
    run = strap.add_run("INTELLIGENT INDUSTRIAL ALARM · VALIDATION BRIEF")
    set_run_font(run, size=9, bold=True, color=BLUE)

    title = lines[0].lstrip("# ").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(13)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=28, bold=True, color=NAVY)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(18)
    run = rule.add_run("━" * 18)
    set_run_font(run, size=9, color=BLUE)

    idx = 1
    while idx < len(lines) and not lines[idx].startswith("## "):
        text = lines[idx].strip()
        if text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, text.replace("  ", ""), size=10.5)
        idx += 1

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    set_table_geometry(box, [7600], indent_dxa=220)
    set_repeat_table_header(box.rows[0])
    cell = box.cell(0, 0)
    set_cell_width(cell, 7600)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=160, start=220, bottom=160, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("结论导向：适配层可以降低部分数据集的阈值错配，但不能替代工况建模、动态残差与显式拒绝。")
    set_run_font(run, size=10.5, bold=True, color=NAVY)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(14)
    run = note.add_run("真实公开数据工程迁移验证 · 三数据集 · 三随机种子 · 点级/事件级/不确定性联合报告")
    set_run_font(run, size=9.5, color="667085")
    doc.add_page_break()
    return idx


def clean_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = clean_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def column_widths(rows: list[list[str]], total: int = 9360) -> list[int]:
    cols = max(len(row) for row in rows)
    weights = []
    for col in range(cols):
        lengths = [len(re.sub(r"\*\*|`", "", row[col])) if col < len(row) else 0 for row in rows]
        weights.append(max(7, min(30, max(lengths, default=7))))
    denom = sum(weights)
    widths = [max(720, round(total * w / denom)) for w in weights]
    delta = total - sum(widths)
    widths[-1] += delta
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    widths = column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        set_cant_split(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=7.6 if cols >= 6 else 8.1)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, heading: str) -> None:
    if heading not in FIGURES:
        return
    path, caption = FIGURES[heading]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(path), width=Inches(6.15))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", heading)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.keep_with_next = False


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [9000], indent_dxa=160)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9000)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code_lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor.from_string("24292F")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_source_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    run = p.add_run(
        "证据说明：所有性能值来自仓库中冻结 JSON 机器报告；书籍与论文 DOI 用于方法溯源。"
        "本报告不把合成 smoke、代理协议或选择性子集结果表述为论文原始分数。"
    )
    set_run_font(run, size=8.5, color="667085")


def build() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    idx = add_cover(doc, lines)
    in_code = False
    code_lines: list[str] = []
    ordered_num_id: int | None = None

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            ordered_num_id = None
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue

        if stripped.startswith("## "):
            ordered_num_id = None
            heading = stripped[3:].strip()
            if heading in PAGE_BREAK_BEFORE:
                doc.add_page_break()
            p = doc.add_paragraph(heading, style="Heading 1")
            p.paragraph_format.keep_with_next = True
            add_figure(doc, heading)
            idx += 1
            continue
        if stripped.startswith("### "):
            ordered_num_id = None
            heading = stripped[4:].strip()
            p = doc.add_paragraph(heading, style="Heading 2")
            p.paragraph_format.keep_with_next = True
            add_figure(doc, heading)
            idx += 1
            continue
        if stripped.startswith("#### "):
            ordered_num_id = None
            doc.add_paragraph(stripped[5:].strip(), style="Heading 3")
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and is_table_separator(lines[idx + 1]):
            ordered_num_id = None
            rows = [clean_table_row(stripped)]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(clean_table_row(lines[idx]))
                idx += 1
            add_table(doc, rows)
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            if ordered_num_id is None:
                ordered_num_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, ordered_num_id)
            add_inline(p, ordered.group(2))
            idx += 1
            continue
        if stripped.startswith("- "):
            ordered_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            idx += 1
            continue

        ordered_num_id = None
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        p.paragraph_format.keep_together = len(stripped) < 240
        add_inline(p, stripped.replace("  ", ""))
        idx += 1

    properties = doc.core_properties
    properties.title = "IIA Benchmark 单变量与多变量报警验证、分布分析及适应性改进报告"
    properties.subject = "单变量、多变量算法验证与跨数据集适配"
    properties.author = "IIA Benchmark Engineering Team"
    properties.keywords = "industrial alarm, univariate, multivariate, distribution shift, adaptation"
    properties.comments = "Generated from frozen repository evidence on 2026-09-03."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
